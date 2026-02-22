"""
Sprint 11 — BLOQUE E: Decision Report Exports (DOCX + XLSX).
Builds downloadable reports from DecisionPack + DecisionRanking + DecisionOutcome.
"""
import io
from datetime import datetime
from typing import Optional
from uuid import UUID

from sqlalchemy.orm import Session

from backend.src.database.models import (
    DecisionOutcome,
    DecisionPack,
    DecisionRanking,
)


class ReportBuilder:
    """Builds DOCX and XLSX reports for a given decision."""

    def __init__(self, db: Session):
        self.db = db

    def _load_decision(self, decision_id: UUID, org_ad_account_ids: list):
        """Load decision pack, ranking, and outcomes. Returns (pack, ranking, outcomes)."""
        pack = self.db.query(DecisionPack).filter(
            DecisionPack.id == decision_id,
            DecisionPack.ad_account_id.in_(org_ad_account_ids),
        ).first()
        if not pack:
            return None, None, []

        ranking = self.db.query(DecisionRanking).filter(
            DecisionRanking.decision_id == decision_id,
        ).first()

        outcomes = self.db.query(DecisionOutcome).filter(
            DecisionOutcome.decision_id == decision_id,
        ).all()

        return pack, ranking, outcomes

    def build_docx(self, decision_id: UUID, org_ad_account_ids: list) -> Optional[io.BytesIO]:
        """Build a DOCX report for a decision. Returns BytesIO or None if not found."""
        from docx import Document
        from docx.shared import Pt, Inches

        pack, ranking, outcomes = self._load_decision(decision_id, org_ad_account_ids)
        if not pack:
            return None

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(10)

        # Title
        doc.add_heading(f"Decision Report: {pack.entity_name or pack.entity_id}", level=1)
        doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}")

        # Executive Summary
        doc.add_heading("Executive Summary", level=2)
        action_desc = (pack.action_type.value if pack.action_type else "action").replace("_", " ")
        entity_desc = pack.entity_name or pack.entity_id or "entity"
        risk_level = "low" if (pack.risk_score or 0) < 0.3 else "medium" if (pack.risk_score or 0) < 0.7 else "high"
        doc.add_paragraph(
            f"This report documents a {action_desc} decision for {entity_desc}. "
            f"The decision was assessed at {risk_level} risk "
            f"(score: {f'{pack.risk_score:.2f}' if pack.risk_score else 'N/A'}) "
            f"and was created on {pack.created_at.strftime('%B %d, %Y') if pack.created_at else 'unknown date'}. "
            f"Source: {pack.source or 'manual'}."
        )

        # Summary table
        doc.add_heading("Summary", level=2)
        table = doc.add_table(rows=7, cols=2)
        table.style = "Light Grid Accent 1"
        summary_data = [
            ("Decision ID", str(pack.id)),
            ("State", pack.state.value if pack.state else "unknown"),
            ("Action Type", pack.action_type.value if pack.action_type else "N/A"),
            ("Entity", f"{pack.entity_type} / {pack.entity_id}"),
            ("Risk Score", f"{pack.risk_score:.2f}" if pack.risk_score else "N/A"),
            ("Source", pack.source or "N/A"),
            ("Created At", pack.created_at.isoformat() if pack.created_at else "N/A"),
        ]
        for i, (key, val) in enumerate(summary_data):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = str(val)

        # Rationale
        doc.add_heading("Rationale", level=2)
        doc.add_paragraph(pack.rationale or "No rationale provided.")

        # Before/After
        doc.add_heading("Before Snapshot", level=2)
        if pack.before_snapshot:
            for k, v in pack.before_snapshot.items():
                doc.add_paragraph(f"{k}: {v}", style="List Bullet")
        else:
            doc.add_paragraph("No before snapshot available.")

        doc.add_heading("After Proposal", level=2)
        if pack.after_proposal:
            for k, v in pack.after_proposal.items():
                doc.add_paragraph(f"{k}: {v}", style="List Bullet")
        else:
            doc.add_paragraph("No after proposal available.")

        # Metrics Comparison
        if pack.before_snapshot and pack.after_proposal:
            doc.add_heading("Metrics Comparison", level=2)
            all_keys = sorted(set(pack.before_snapshot.keys()) | set(pack.after_proposal.keys()))
            metrics_table = doc.add_table(rows=1, cols=3)
            metrics_table.style = "Light Grid Accent 1"
            metrics_table.rows[0].cells[0].text = "Metric"
            metrics_table.rows[0].cells[1].text = "Before"
            metrics_table.rows[0].cells[2].text = "After"
            for key in all_keys:
                row = metrics_table.add_row()
                row.cells[0].text = key
                row.cells[1].text = str(pack.before_snapshot.get(key, "—"))
                row.cells[2].text = str(pack.after_proposal.get(key, "—"))

        # Policy Checks
        doc.add_heading("Policy Checks", level=2)
        if pack.policy_checks:
            for check in pack.policy_checks:
                if isinstance(check, dict):
                    name = check.get("policy", check.get("name", "Unknown"))
                    passed = check.get("passed", check.get("result", "N/A"))
                    doc.add_paragraph(f"{name}: {passed}", style="List Bullet")
        else:
            doc.add_paragraph("No policy checks recorded.")

        # Ranking Scores
        if ranking:
            doc.add_heading("Ranking Scores", level=2)
            scores_table = doc.add_table(rows=5, cols=2)
            scores_table.style = "Light Grid Accent 1"
            scores = [
                ("Total", f"{ranking.score_total:.2f}"),
                ("Impact", f"{ranking.score_impact:.2f}"),
                ("Risk", f"{ranking.score_risk:.2f}"),
                ("Confidence", f"{ranking.score_confidence:.2f}"),
                ("Freshness", f"{ranking.score_freshness:.2f}"),
            ]
            for i, (key, val) in enumerate(scores):
                scores_table.rows[i].cells[0].text = key
                scores_table.rows[i].cells[1].text = val

        # Outcomes
        if outcomes:
            doc.add_heading("Outcomes", level=2)
            for outcome in outcomes:
                doc.add_paragraph(
                    f"Horizon: {outcome.horizon_minutes}min | "
                    f"Label: {outcome.outcome_label.value if outcome.outcome_label else 'unknown'} | "
                    f"Confidence: {outcome.confidence:.2f}",
                    style="List Bullet",
                )

        # Execution Result
        doc.add_heading("Execution Result", level=2)
        if pack.execution_result:
            for k, v in pack.execution_result.items():
                doc.add_paragraph(f"{k}: {v}", style="List Bullet")
        else:
            doc.add_paragraph("Not yet executed.")

        # Recommendations
        doc.add_heading("Recommendations", level=2)
        if outcomes:
            positive = [o for o in outcomes if o.outcome_label and o.outcome_label.value == "positive"]
            if len(positive) > len(outcomes) / 2:
                doc.add_paragraph(
                    "Based on outcome data, this type of decision shows positive results. "
                    "Consider scaling similar actions."
                )
            else:
                doc.add_paragraph(
                    "Outcome data suggests mixed results. "
                    "Review individual outcome metrics before scaling."
                )
        else:
            doc.add_paragraph(
                "No outcome data available yet. "
                "Schedule outcome capture to measure impact."
            )

        # Data Sources
        doc.add_heading("Data Sources", level=2)
        doc.add_paragraph(f"Decision source: {pack.source or 'Manual'}", style="List Bullet")
        doc.add_paragraph(f"Report generated: {datetime.utcnow().isoformat()} UTC", style="List Bullet")
        doc.add_paragraph("Data from: Meta Ads API via Meta Ops Agent", style="List Bullet")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def build_xlsx(self, decision_id: UUID, org_ad_account_ids: list) -> Optional[io.BytesIO]:
        """Build an XLSX report for a decision. Returns BytesIO or None if not found."""
        from openpyxl import Workbook

        pack, ranking, outcomes = self._load_decision(decision_id, org_ad_account_ids)
        if not pack:
            return None

        wb = Workbook()

        # Sheet 1: Summary
        ws = wb.active
        ws.title = "Summary"
        ws.append(["Field", "Value"])
        ws.append(["Decision ID", str(pack.id)])
        ws.append(["State", pack.state.value if pack.state else "unknown"])
        ws.append(["Action Type", pack.action_type.value if pack.action_type else "N/A"])
        ws.append(["Entity Type", pack.entity_type or ""])
        ws.append(["Entity ID", pack.entity_id or ""])
        ws.append(["Entity Name", pack.entity_name or ""])
        ws.append(["Risk Score", pack.risk_score or 0])
        ws.append(["Source", pack.source or ""])
        ws.append(["Rationale", pack.rationale or ""])
        ws.append(["Created At", pack.created_at.isoformat() if pack.created_at else ""])
        ws.append(["Executed At", pack.executed_at.isoformat() if pack.executed_at else ""])

        # Before/After
        if pack.before_snapshot:
            ws.append([])
            ws.append(["Before Snapshot"])
            for k, v in pack.before_snapshot.items():
                ws.append([k, str(v)])

        if pack.after_proposal:
            ws.append([])
            ws.append(["After Proposal"])
            for k, v in pack.after_proposal.items():
                ws.append([k, str(v)])

        # Sheet 2: Policy Checks
        ws2 = wb.create_sheet("Policy Checks")
        ws2.append(["Policy", "Passed", "Details"])
        if pack.policy_checks:
            for check in pack.policy_checks:
                if isinstance(check, dict):
                    ws2.append([
                        check.get("policy", check.get("name", "")),
                        str(check.get("passed", check.get("result", ""))),
                        check.get("message", check.get("detail", "")),
                    ])

        # Sheet 3: Ranking
        ws3 = wb.create_sheet("Ranking")
        ws3.append(["Score", "Value"])
        if ranking:
            ws3.append(["Total", ranking.score_total])
            ws3.append(["Impact", ranking.score_impact])
            ws3.append(["Risk", ranking.score_risk])
            ws3.append(["Confidence", ranking.score_confidence])
            ws3.append(["Freshness", ranking.score_freshness])
            ws3.append(["Rank Version", ranking.rank_version])
            if ranking.explanation_json:
                ws3.append([])
                ws3.append(["Explanation"])
                for k, v in ranking.explanation_json.items():
                    ws3.append([k, str(v)])

        # Sheet 4: Outcomes
        ws4 = wb.create_sheet("Outcomes")
        ws4.append(["Horizon (min)", "Label", "Confidence", "Dry Run", "Executed At", "Notes"])
        for outcome in outcomes:
            ws4.append([
                outcome.horizon_minutes,
                outcome.outcome_label.value if outcome.outcome_label else "unknown",
                outcome.confidence,
                outcome.dry_run,
                outcome.executed_at.isoformat() if outcome.executed_at else "",
                outcome.notes or "",
            ])

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf
